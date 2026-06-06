#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tao 3 so do Use Case cho he thong AutoHT bang PlantUML
Output: UC1_KhachHang.puml/.png, UC2_NhanVien.puml/.png, UC3_Admin.puml/.png
        UseCase_AutoHT.docx (A4 landscape, 1 so do / trang)
"""

import os, sys, zlib
sys.stdout = __import__('io').TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

BASE = os.path.dirname(os.path.abspath(__file__))

try:
    import requests
    HAS_REQ = True
except ImportError:
    HAS_REQ = False
    print("[WARN] requests chua cai. Chay: pip install requests")

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx chua cai. Chay: pip install python-docx")


# ── PlantUML Web Server Encoder ──────────────────────────────────────────────

def _e6(b):
    if b < 10:  return chr(48 + b)
    b -= 10
    if b < 26:  return chr(65 + b)
    b -= 26
    if b < 26:  return chr(97 + b)
    b -= 26
    return '-' if b == 0 else '_'

def puml_encode(text):
    raw = zlib.compress(text.encode('utf-8'), 9)[2:-4]
    out = []
    for i in range(0, len(raw), 3):
        ch = raw[i:i+3]
        b1, b2, b3 = ch[0], ch[1] if len(ch)>1 else 0, ch[2] if len(ch)>2 else 0
        s = _e6(b1>>2) + _e6(((b1&3)<<4)|(b2>>4)) + _e6(((b2&15)<<2)|(b3>>6)) + _e6(b3&63)
        out.append(s if len(ch)==3 else s[:len(ch)+1])
    return ''.join(out)

def render_png(puml_text, out_path):
    if not HAS_REQ:
        print(f"[SKIP] {out_path} (requests chua cai)")
        return False
    try:
        url = "https://www.plantuml.com/plantuml/png/" + puml_encode(puml_text)
        r = requests.get(url, timeout=30)
        r.raise_for_status()
        with open(out_path, 'wb') as f:
            f.write(r.content)
        print(f"[OK] PNG: {out_path}")
        return True
    except Exception as e:
        print(f"[WARN] Khong render duoc {out_path}: {e}")
        return False


# ── 3 PlantUML Diagrams ──────────────────────────────────────────────────────

PUML_1 = """\
@startuml UC1_KhachHang
title AutoHT -- So do Use Case 1/3: Phan he Khach Hang (Client System)

skinparam defaultFontName Arial
skinparam defaultFontSize 13
skinparam Nodesep 35
skinparam Ranksep 60

skinparam actor {
  BackgroundColor #FDFEFE
  BorderColor #1A5276
  FontSize 14
  FontStyle bold
}
skinparam usecase {
  BackgroundColor #D6EAF8
  BorderColor #1A5276
  FontSize 13
}
skinparam package {
  BackgroundColor #EBF5FB
  BorderColor #2874A6
  FontSize 14
  FontStyle bold
}
skinparam rectangle {
  BackgroundColor #F8FBFF
  BorderColor #1A5276
  BorderThickness 2
  FontSize 15
  FontStyle bold
}
skinparam ArrowColor #1A5276
skinparam ArrowFontSize 12

actor "Khach Hang" as KH

rectangle "He Thong AutoHT -- Phan he Khach Hang" {

  package "Quan ly Tai Khoan" {
    (Dang ky tai khoan) as UC1
    (Dang nhap he thong) as UC2
    (Xac thuc Gmail / OTP) as UC3
    (Quan ly ho so ca nhan) as UC4
  }

  package "Tra cuu & Tuong tac Xe" {
    (Tim kiem & So sanh xe) as UC5
    (Xem mo hinh xe 3D) as UC6
  }

  package "Giao dich Truc tuyen" {
    (Dat mua xe truc tuyen) as UC7
    (Dat thue xe truc tuyen) as UC8
    (Theo doi lich su don hang) as UC9
    note as N1
      Yeu cau: Dang nhap
      truoc khi giao dich
    end note
  }

}

KH -right-> UC1
KH -right-> UC2
KH -right-> UC4
KH -right-> UC5
KH -right-> UC7
KH -right-> UC8
KH -right-> UC9

UC1 ..> UC3 : <<include>>
UC5 ..> UC6 : <<extend>>

@enduml
"""

PUML_2 = """\
@startuml UC2_NhanVien
title AutoHT -- So do Use Case 2/3: Phan he Nghiep vu Noi bo (Staff System)

left to right direction

skinparam defaultFontName Arial
skinparam defaultFontSize 12

skinparam actor {
  BackgroundColor #FDFEFE
  BorderColor #1E8449
  FontSize 13
  FontStyle bold
}
skinparam usecase {
  BackgroundColor #D5F5E3
  BorderColor #1E8449
  FontSize 12
}
skinparam package {
  BackgroundColor #EAFAF1
  BorderColor #27AE60
  FontSize 13
  FontStyle bold
}
skinparam rectangle {
  BackgroundColor #F9FEF9
  BorderColor #1E8449
  BorderThickness 2
  FontSize 15
  FontStyle bold
}
skinparam ArrowColor #1E8449
skinparam ArrowFontSize 11
skinparam Nodesep 35
skinparam Ranksep 50

actor "NV Kho (Warehouse)" as NVKho
actor "NV Kinh Doanh (Sales)" as NVKD
actor "NV Ky Thuat (Technical)" as NVKT
actor "Ke Toan (Accounting)" as KeToan

rectangle "He Thong AutoHT -- Phan he Nghiep vu Noi bo" {

  package "Quan ly Kho" {
    (Nhap kho & Kiem dinh PDI) as UC_K1
    (Quan ly xe don vi - XeDonVi) as UC_K2
    (Kiem ke & Thong ke ton kho) as UC_K3
  }

  package "Kinh Doanh & CRM" {
    (Quan ly thong tin khach hang) as UC_D1
    (Tao don & Tinh phuong an tra gop) as UC_D2
    (Phe duyet & Xu ly giao dich) as UC_D3
  }

  package "Ky Thuat & Hau Mai" {
    (Tiep nhan lich bao duong) as UC_T1
    (Lap phieu & Quan ly sua chua) as UC_T2
    (Dieu phoi kho phu tung) as UC_T3
  }

  package "Tai Chinh & Ke Toan" {
    (Quan ly thu chi & Hoa don) as UC_C1
    (Lap bao cao tai chinh) as UC_C2
  }

}

NVKho --> UC_K1
NVKho --> UC_K2
NVKho --> UC_K3

NVKD --> UC_D1
NVKD --> UC_D2
NVKD --> UC_D3

NVKT --> UC_T1
NVKT --> UC_T2
NVKT --> UC_T3

KeToan --> UC_C1
KeToan --> UC_C2

UC_K1 ..> UC_K2 : <<include>>
UC_D2 ..> UC_D3 : <<include>>
UC_T1 ..> UC_T2 : <<include>>
UC_T2 ..> UC_T3 : <<include>>

@enduml
"""

PUML_3 = """\
@startuml UC3_Admin
title AutoHT -- So do Use Case 3/3: Phan he Quan tri He thong (Admin System)

skinparam defaultFontName Arial
skinparam defaultFontSize 13
skinparam Nodesep 35
skinparam Ranksep 60

skinparam actor {
  BackgroundColor #FDFEFE
  BorderColor #922B21
  FontSize 14
  FontStyle bold
}
skinparam usecase {
  BackgroundColor #FADBD8
  BorderColor #922B21
  FontSize 13
}
skinparam package {
  BackgroundColor #FDEDEC
  BorderColor #C0392B
  FontSize 14
  FontStyle bold
}
skinparam rectangle {
  BackgroundColor #FFF9F9
  BorderColor #922B21
  BorderThickness 2
  FontSize 15
  FontStyle bold
}
skinparam ArrowColor #922B21
skinparam ArrowFontSize 12

actor "Quan Tri Vien (Admin)" as Admin

rectangle "He Thong AutoHT -- Phan he Quan tri" {

  package "Quan ly Tai Khoan & Phan Quyen" {
    (Tao / Sua / Xoa tai khoan NV) as UC_Q1
    (Phan quyen theo vai tro - RBAC) as UC_Q2
    (Khoa & Mo khoa tai khoan) as UC_Q3
  }

  package "Cau hinh He thong" {
    (Quan ly danh muc xe & Hang xe) as UC_H1
    (Cau hinh tham so & Chinh sach) as UC_H2
  }

  package "Giam sat & Bao cao Tong quan" {
    (Xem Dashboard & KPI tong quan) as UC_B1
    (Xem bao cao Doanh thu & Kho) as UC_B2
    (Tra cuu Nhat ky he thong) as UC_B3
  }

}

Admin -right-> UC_Q1
Admin -right-> UC_Q2
Admin -right-> UC_Q3
Admin -right-> UC_H1
Admin -right-> UC_H2
Admin -right-> UC_B1
Admin -right-> UC_B2
Admin -right-> UC_B3

UC_Q1 ..> UC_Q2 : <<include>>
UC_B1 ..> UC_B2 : <<extend>>

@enduml
"""

PUML_TONG_QUAT = """\
@startuml UC_TongQuat
title AutoHT -- So do Use Case Tong quat (Overview)

left to right direction

skinparam defaultFontName Arial
skinparam defaultFontSize 13
skinparam Nodesep 25
skinparam Ranksep 80

skinparam actor {
  BackgroundColor #FDFEFE
  BorderColor #2C3E50
  FontSize 13
  FontStyle bold
}
skinparam usecase {
  BackgroundColor #EBF5FB
  BorderColor #2874A6
  FontSize 12
}
skinparam package {
  FontSize 13
  FontStyle bold
}
skinparam rectangle {
  BackgroundColor #FAFAFA
  BorderColor #2C3E50
  BorderThickness 2
  FontSize 15
  FontStyle bold
}
skinparam ArrowColor #2C3E50
skinparam ArrowFontSize 11

' ── Actors (declared top-to-bottom = position top-to-bottom) ────
actor "Khach Hang" as KH
actor "NV Kinh Doanh" as NVKD
actor "NV Kho" as NVKho
actor "NV Ky Thuat" as NVKT
actor "Ke Toan" as KeToan
actor "Admin" as Admin

' ── System Boundary ─────────────────────────────────────────────
rectangle "He Thong AutoHT" {

  package "Phan he Khach Hang" #EBF5FB {
    (Quan ly Tai Khoan & Dang nhap) as UC1
    (Tra cuu, So sanh & Xe 3D) as UC2
    (Giao dich Truc tuyen) as UC3
  }

  package "Kinh Doanh & CRM" #EAFAF1 {
    (Kinh Doanh & Quan ly Khach hang) as UC4
  }

  package "Quan ly Kho" #EAFAF1 {
    (Quan ly Kho & Kiem dinh PDI) as UC5
  }

  package "Ky Thuat & Hau Mai" #EAFAF1 {
    (Ky Thuat & Bao duong) as UC6
  }

  package "Tai Chinh & Ke Toan" #EAFAF1 {
    (Quan ly Thu chi & Bao cao) as UC7
  }

  package "Phan he Quan tri" #FDEDEC {
    (Quan tri Tai Khoan & Phan Quyen) as UC8
    (Cau hinh & Giam sat He thong) as UC9
  }

}

' ── Connections ─────────────────────────────────────────────────
KH --> UC1
KH --> UC2
KH --> UC3

NVKD --> UC4
NVKho --> UC5
NVKT --> UC6
KeToan --> UC7

Admin --> UC8
Admin --> UC9

' ── Key Relationships ────────────────────────────────────────────
UC3 ..> UC1 : <<include>>

@enduml
"""

DIAGRAMS = [
    {
        "name":  "UC1_KhachHang",
        "title": "So do Use Case 1/3 — Phan he Khach Hang (Client System)",
        "puml":  PUML_1,
    },
    {
        "name":  "UC2_NhanVien",
        "title": "So do Use Case 2/3 — Phan he Nghiep vu Noi bo (Staff System)",
        "puml":  PUML_2,
    },
    {
        "name":  "UC3_Admin",
        "title": "So do Use Case 3/3 — Phan he Quan tri He thong (Admin System)",
        "puml":  PUML_3,
    },
    {
        "name":  "UC_TongQuat",
        "title": "So do Use Case Tong quat — He Thong AutoHT",
        "puml":  PUML_TONG_QUAT,
    },
]


# ── Word Document Builder ────────────────────────────────────────────────────

def build_word(results, out_path):
    if not HAS_DOCX:
        print("[SKIP] Word (python-docx chua cai)")
        return
    doc = Document()

    for i, item in enumerate(results):
        if i > 0:
            doc.add_page_break()

        sec = doc.sections[0] if i == 0 else doc.add_section()
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width  = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.left_margin = sec.right_margin = Cm(1.5)
        sec.top_margin  = sec.bottom_margin = Cm(1.2)

        # Heading
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run(item["title"])
        r.bold = True
        r.font.size = Pt(14)

        # Image
        png = os.path.join(BASE, item["name"] + ".png")
        if os.path.exists(png):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(png, width=Inches(10.2))
        else:
            doc.add_paragraph(f"[Hinh anh chua render: {png}]")
            doc.add_paragraph(f"Hay mo file {item['name']}.puml tai plantuml.com de render.")

    doc.save(out_path)
    print(f"[OK] Word: {out_path}")


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    results = []
    for d in DIAGRAMS:
        # Save .puml
        puml_path = os.path.join(BASE, d["name"] + ".puml")
        with open(puml_path, 'w', encoding='utf-8') as f:
            f.write(d["puml"])
        print(f"[OK] PUML: {puml_path}")

        # Render PNG
        png_path = os.path.join(BASE, d["name"] + ".png")
        render_png(d["puml"], png_path)

        results.append(d)

    # Build Word
    build_word(results, os.path.join(BASE, "UseCase_AutoHT.docx"))

    print("\n=== Hoan thanh! ===")
    print("File da tao:")
    for d in DIAGRAMS:
        print(f"  {d['name']}.puml  — dan vao https://www.plantuml.com/plantuml/ de render")
        print(f"  {d['name']}.png   — anh da render (neu co internet)")
    print("  UseCase_AutoHT.docx — Word A4 ngang chua ca 3 so do")

if __name__ == '__main__':
    main()
