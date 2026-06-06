"""
Tao so do ERD cho du an AutoHT - Phien ban chinh sua
Layout gon gang, khong cheo duong, nhan tieng Viet chuan.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ═══════════════════════════════════════════════════════════════
# THIET LAP CANVAS  (compact: 25 x 17)
# ═══════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(25, 17))
ax.set_xlim(0, 25)
ax.set_ylim(0, 17)
ax.axis('off')
fig.patch.set_facecolor('#EEF2F7')
ax.set_facecolor('#EEF2F7')

# ── MAU SAC ──
COL = {
    'MauXe':              '#1B4332',
    'XeDonVi':            '#7B341E',
    'DonHang':            '#553C9A',
    'KhachHang':          '#1A365D',
    'PhieuDichVu':        '#2C5282',
    'ChiTietPhieuDichVu': '#234E52',
    'LinhKienPhuTung':    '#702459',
}
PK_BG   = '#FFFBEA'
FK_BG   = '#EBF8FF'
PKFK_BG = '#F0FFF4'
WHITE   = '#FFFFFF'
BORDER  = '#2D3748'
LINE_C  = '#4A5568'

# ═══════════════════════════════════════════════════════════════
# HAM VE ENTITY BOX
# ═══════════════════════════════════════════════════════════════
ENTITY_W = 5.0
HEAD_H   = 0.78
ROW_H    = 0.52
PAD      = 0.10

def entity_height(n_attrs):
    return HEAD_H + n_attrs * ROW_H + PAD

def draw_entity(ax, x, y_top, name, display, attrs):
    """Ve entity box. Tra ve (bot_y, mid_y)."""
    w   = ENTITY_W
    col = COL[name]
    th  = entity_height(len(attrs))

    # Shadow
    shadow = FancyBboxPatch((x+0.10, y_top-th-0.10), w, th,
        boxstyle='round,pad=0.04', facecolor='#00000025',
        edgecolor='none', zorder=1)
    ax.add_patch(shadow)

    # Main box
    box = FancyBboxPatch((x, y_top-th), w, th,
        boxstyle='round,pad=0.04', facecolor=WHITE,
        edgecolor=BORDER, linewidth=2.0, zorder=2)
    ax.add_patch(box)

    # Header
    hdr = FancyBboxPatch((x+0.04, y_top-HEAD_H+0.03), w-0.08, HEAD_H-0.04,
        boxstyle='round,pad=0.03', facecolor=col,
        edgecolor='none', zorder=3)
    ax.add_patch(hdr)

    # Header text
    lines = display.split('\n')
    if len(lines) == 2:
        ax.text(x+w/2, y_top-HEAD_H/2+0.11, lines[0],
            ha='center', va='center', fontsize=9.5, fontweight='bold',
            color='white', zorder=5)
        ax.text(x+w/2, y_top-HEAD_H/2-0.16, lines[1],
            ha='center', va='center', fontsize=8.0,
            color='#FFFFFFCC', zorder=5)
    else:
        ax.text(x+w/2, y_top-HEAD_H/2, display,
            ha='center', va='center', fontsize=9.5, fontweight='bold',
            color='white', zorder=5)

    # Duong ke ngang sau header
    ax.plot([x, x+w], [y_top-HEAD_H, y_top-HEAD_H],
        color=BORDER, linewidth=1.2, zorder=4)

    # Attributes
    for i, attr in enumerate(attrs):
        row_top = y_top - HEAD_H - i*ROW_H
        row_bot = row_top - ROW_H

        # Nen mau
        is_pk = attr.get('pk', False)
        is_fk = attr.get('fk', False)
        if is_pk and is_fk: bg = PKFK_BG
        elif is_pk:          bg = PK_BG
        elif is_fk:          bg = FK_BG
        else:                bg = WHITE

        rect = plt.Rectangle((x, row_bot), w, ROW_H,
            facecolor=bg, edgecolor='none', zorder=3)
        ax.add_patch(rect)

        # Duong ke ngang giua cac row
        if i > 0:
            ax.plot([x+0.12, x+w-0.12], [row_top, row_top],
                color='#CBD5E0', linewidth=0.7, zorder=4)

        row_mid = (row_top + row_bot) / 2

        # Tag [PK] / [FK]
        if is_pk and is_fk:
            tag = '[PK+FK]'
            tcolor = '#276749'
        elif is_pk:
            tag = '[PK]   '
            tcolor = '#975A16'
        elif is_fk:
            tag = '[FK]   '
            tcolor = '#2B6CB0'
        else:
            tag = '       '
            tcolor = '#A0AEC0'

        ax.text(x+0.15, row_mid, tag,
            ha='left', va='center', fontsize=7.0,
            color=tcolor, family='monospace', zorder=5)

        ax.text(x+1.02, row_mid, attr['col'],
            ha='left', va='center', fontsize=8.0,
            color='#1A202C',
            fontweight='bold' if (is_pk or is_fk) else 'normal',
            zorder=5)

        ax.text(x+w-0.12, row_mid, attr['dtype'],
            ha='right', va='center', fontsize=7.0,
            color='#718096', style='italic', zorder=5)

    bot_y = y_top - th
    mid_y = (y_top + bot_y) / 2
    return bot_y, mid_y

# ═══════════════════════════════════════════════════════════════
# DINH NGHIA ENTITIES
# ═══════════════════════════════════════════════════════════════
ROW1_YTOP = 15.8   # Hang 1
ROW2_YTOP =  8.9   # Hang 2 (PhieuDichVu)
ROW3_YTOP =  3.6   # Hang 3 (ChiTiet, LinhKien)

# X toa do hang 1 (cach nhau 1.2 don vi)
X = {
    'MauXe':              0.5,
    'XeDonVi':            6.7,
    'DonHang':           12.9,
    'KhachHang':         19.1,
    'PhieuDichVu':        6.7,   # thang hang voi XeDonVi
    'ChiTietPhieuDichVu': 2.5,
    'LinhKienPhuTung':   13.5,
}

ENTITIES = {
    'MauXe': {
        'display': 'MauXe\n(Mau Xe)',
        'ytop': ROW1_YTOP,
        'attrs': [
            {'col':'ModelID',   'dtype':'INT PK Identity', 'pk':True},
            {'col':'ModelName', 'dtype':'NVARCHAR(100)'},
            {'col':'Brand',     'dtype':'NVARCHAR(50)'},
            {'col':'BasePrice', 'dtype':'DECIMAL(18,2)'},
            {'col':'Engine',    'dtype':'NVARCHAR(50)'},
        ]
    },
    'XeDonVi': {
        'display': 'XeDonVi\n(Xe Don Vi)',
        'ytop': ROW1_YTOP,
        'attrs': [
            {'col':'VIN',      'dtype':'VARCHAR(17) PK',   'pk':True},
            {'col':'ModelID',  'dtype':'INT FK',            'fk':True},
            {'col':'Color',    'dtype':'NVARCHAR(30)'},
            {'col':'PDIStatus','dtype':'NVARCHAR(50)'},
            {'col':'Status',   'dtype':'NVARCHAR(50)'},
        ]
    },
    'DonHang': {
        'display': 'DonHang\n(Don Hang)',
        'ytop': ROW1_YTOP,
        'attrs': [
            {'col':'OrderID',    'dtype':'INT PK Identity', 'pk':True},
            {'col':'CustomerID', 'dtype':'INT FK',           'fk':True},
            {'col':'VIN',        'dtype':'VARCHAR(17) FK',   'fk':True},
            {'col':'OrderType',  'dtype':'NVARCHAR(20)'},
            {'col':'TotalAmount','dtype':'DECIMAL(18,2)'},
            {'col':'OrderStatus','dtype':'NVARCHAR(50)'},
        ]
    },
    'KhachHang': {
        'display': 'KhachHang\n(Khach Hang)',
        'ytop': ROW1_YTOP,
        'attrs': [
            {'col':'CustomerID','dtype':'INT PK Identity', 'pk':True},
            {'col':'FullName',  'dtype':'NVARCHAR(100)'},
            {'col':'Email',     'dtype':'VARCHAR(100)'},
            {'col':'Phone',     'dtype':'VARCHAR(15)'},
            {'col':'Address',   'dtype':'NVARCHAR(255)'},
        ]
    },
    'PhieuDichVu': {
        'display': 'PhieuDichVu\n(Phieu Dich Vu)',
        'ytop': ROW2_YTOP,
        'attrs': [
            {'col':'TicketID',   'dtype':'INT PK Identity', 'pk':True},
            {'col':'VIN',        'dtype':'VARCHAR(17) FK',  'fk':True},
            {'col':'CustomerID', 'dtype':'INT FK',           'fk':True},
            {'col':'CreateDate', 'dtype':'DATETIME'},
            {'col':'Description','dtype':'NVARCHAR(MAX)'},
            {'col':'TotalCost',  'dtype':'DECIMAL(18,2)'},
        ]
    },
    'ChiTietPhieuDichVu': {
        'display': 'ChiTietPhieuDichVu\n(Chi Tiet Phieu DV)',
        'ytop': ROW3_YTOP,
        'attrs': [
            {'col':'TicketID','dtype':'INT PK,FK', 'pk':True,'fk':True},
            {'col':'PartID',  'dtype':'INT PK,FK', 'pk':True,'fk':True},
            {'col':'Quantity','dtype':'INT'},
        ]
    },
    'LinhKienPhuTung': {
        'display': 'LinhKienPhuTung\n(Linh Kien Phu Tung)',
        'ytop': ROW3_YTOP,
        'attrs': [
            {'col':'PartID',       'dtype':'INT PK Identity','pk':True},
            {'col':'PartName',     'dtype':'NVARCHAR(100)'},
            {'col':'Price',        'dtype':'DECIMAL(18,2)'},
            {'col':'StockQuantity','dtype':'INT'},
        ]
    },
}

# ═══════════════════════════════════════════════════════════════
# VE TAT CA ENTITIES — luu lai toa do
# ═══════════════════════════════════════════════════════════════
B = {}  # bounds: name -> {x, y_top, bot, mid_y, cx, right}

for name, edef in ENTITIES.items():
    x     = X[name]
    y_top = edef['ytop']
    bot, mid_y = draw_entity(ax, x, y_top, name, edef['display'], edef['attrs'])
    B[name] = {
        'x':     x,
        'y_top': y_top,
        'bot':   bot,
        'mid_y': mid_y,
        'cx':    x + ENTITY_W/2,
        'right': x + ENTITY_W,
    }

# ═══════════════════════════════════════════════════════════════
# HAM VE DUONG NOI VA NHAN
# ═══════════════════════════════════════════════════════════════
def seg(x1, y1, x2, y2, c=LINE_C, lw=2.0):
    ax.plot([x1, x2], [y1, y2], color=c, linewidth=lw, zorder=6)

def card(x, y, text, bg=LINE_C):
    """Nhan ban so (1 hoac N)."""
    ax.text(x, y, text,
        ha='center', va='center', fontsize=10, fontweight='bold',
        color='white', zorder=12,
        bbox=dict(boxstyle='round,pad=0.30', facecolor=bg,
                  edgecolor='white', linewidth=1.5))

def rel_lbl(x, y, text):
    """Nhan ten moi quan he."""
    ax.text(x, y, text,
        ha='center', va='center', fontsize=8.5,
        color='#2D3748', style='italic', zorder=11,
        bbox=dict(boxstyle='round,pad=0.22', facecolor='white',
                  edgecolor='#A0AEC0', linewidth=1.0, alpha=0.95))

# ═══════════════════════════════════════════════════════════════
# R1: MauXe (1) ──── (N) XeDonVi
#     Duong ngang thang
# ═══════════════════════════════════════════════════════════════
x1, y1 = B['MauXe']['right'],   B['MauXe']['mid_y']
x2, y2 = B['XeDonVi']['x'],     B['XeDonVi']['mid_y']
seg(x1, y1, x2, y2)
card(x1+0.32, y1+0.30, '1')
card(x2-0.32, y2+0.30, 'N')
rel_lbl((x1+x2)/2, y1+0.55, 'Co')

# ═══════════════════════════════════════════════════════════════
# R2: XeDonVi (1) ──── (1) DonHang   [qua VIN]
#     Duong ngang thang
# ═══════════════════════════════════════════════════════════════
x1, y1 = B['XeDonVi']['right'],  B['XeDonVi']['mid_y']
x2, y2 = B['DonHang']['x'],      B['DonHang']['mid_y']
seg(x1, y1, x2, y2)
card(x1+0.32, y1+0.30, '1')
card(x2-0.32, y2+0.30, '1')
rel_lbl((x1+x2)/2, y1+0.55, 'Nam tren (VIN)')

# ═══════════════════════════════════════════════════════════════
# R3: DonHang (N) ──── (1) KhachHang
#     Duong ngang thang
# ═══════════════════════════════════════════════════════════════
x1, y1 = B['DonHang']['right'],   B['DonHang']['mid_y']
x2, y2 = B['KhachHang']['x'],     B['KhachHang']['mid_y']
seg(x1, y1, x2, y2)
card(x1+0.32, y1+0.30, 'N')
card(x2-0.32, y2+0.30, '1')
rel_lbl((x1+x2)/2, y1+0.55, 'Dat don')

# ═══════════════════════════════════════════════════════════════
# R4: XeDonVi (1) ──── (N) PhieuDichVu
#     Duong thang dung (cung truc x = 9.2)
# ═══════════════════════════════════════════════════════════════
cx_xdv = B['XeDonVi']['cx']   # = 9.2
y_bot_xdv = B['XeDonVi']['bot']
y_top_pdv = B['PhieuDichVu']['y_top']

seg(cx_xdv, y_bot_xdv, cx_xdv, y_top_pdv)
card(cx_xdv+0.45, y_bot_xdv-0.30, '1')
card(cx_xdv+0.45, y_top_pdv+0.30, 'N')
rel_lbl(cx_xdv+1.30, (y_bot_xdv+y_top_pdv)/2, 'Khoi tao')

# ═══════════════════════════════════════════════════════════════
# R5: KhachHang (1) ──── (N) PhieuDichVu
#     Duong hinh chu L: xuat phat tu day KhachHang, di xuong
#     roi re trai den canh phai PhieuDichVu
#     Khong cat qua bat ky entity nao.
# ═══════════════════════════════════════════════════════════════
kh_cx   = B['KhachHang']['cx']        # 21.6
kh_bot  = B['KhachHang']['bot']       # ~12.15

pdv_right = B['PhieuDichVu']['right'] # 11.7
pdv_mid_y = B['PhieuDichVu']['mid_y'] # ~6.70

ELBOW_Y  = (kh_bot + B['PhieuDichVu']['y_top']) / 2  # ~10.75 (giua 2 hang)

# Doan 1: KhachHang.bottom xuong elbow_y
seg(kh_cx, kh_bot, kh_cx, ELBOW_Y)
# Doan 2: ngang trai den x = pdv_right
seg(kh_cx, ELBOW_Y, pdv_right, ELBOW_Y)
# Doan 3: xuong den PhieuDichVu.right mid
seg(pdv_right, ELBOW_Y, pdv_right, pdv_mid_y)

# Ban so
card(kh_cx+0.45, kh_bot-0.30, '1')
card(pdv_right+0.45, pdv_mid_y, 'N')
# Nhan tren doan ngang
rel_lbl((kh_cx+pdv_right)/2, ELBOW_Y+0.38, 'Thuoc ve')

# ═══════════════════════════════════════════════════════════════
# R6: PhieuDichVu (1) ──── (N) ChiTietPhieuDichVu
#     Duong cheo ngan
# ═══════════════════════════════════════════════════════════════
pdv_cx  = B['PhieuDichVu']['cx']
pdv_bot = B['PhieuDichVu']['bot']
ct_cx   = B['ChiTietPhieuDichVu']['cx']
ct_top  = B['ChiTietPhieuDichVu']['y_top']

seg(pdv_cx, pdv_bot, ct_cx, ct_top)
card(pdv_cx-0.25, pdv_bot-0.32, '1')
card(ct_cx+0.25, ct_top+0.32, 'N')
rel_lbl((pdv_cx+ct_cx)/2 - 0.8, (pdv_bot+ct_top)/2, 'Bao gom')

# ═══════════════════════════════════════════════════════════════
# R7: ChiTietPhieuDichVu (N) ──── (1) LinhKienPhuTung
#     Duong ngang (gan nhu thang)
# ═══════════════════════════════════════════════════════════════
ct_right  = B['ChiTietPhieuDichVu']['right']
ct_mid_y  = B['ChiTietPhieuDichVu']['mid_y']
lk_left   = B['LinhKienPhuTung']['x']
lk_mid_y  = B['LinhKienPhuTung']['mid_y']

seg(ct_right, ct_mid_y, lk_left, lk_mid_y)
card(ct_right+0.32, ct_mid_y+0.30, 'N')
card(lk_left-0.32, lk_mid_y+0.30, '1')
rel_lbl((ct_right+lk_left)/2, (ct_mid_y+lk_mid_y)/2 + 0.38, 'Su dung')

# ═══════════════════════════════════════════════════════════════
# TIEU DE
# ═══════════════════════════════════════════════════════════════
ax.text(12.5, 16.75,
    'SO DO THUC THE - QUAN HE (ERD)',
    ha='center', va='center', fontsize=17, fontweight='bold',
    color='#1A202C', zorder=10)
ax.text(12.5, 16.30,
    'Du an AutoHT  --  He thong Quan ly & Mua ban Xe O to Truc tuyen',
    ha='center', va='center', fontsize=10, color='#4A5568', zorder=10)

# ═══════════════════════════════════════════════════════════════
# CHU THICH (goc phai duoi, khong chong len entity nao)
# ═══════════════════════════════════════════════════════════════
LGX, LGY = 19.3, 8.7   # goc tren-trai legend
LGW, LGH = 5.4, 5.6

leg_bg = FancyBboxPatch((LGX, LGY-LGH), LGW, LGH,
    boxstyle='round,pad=0.10', facecolor='white',
    edgecolor=BORDER, linewidth=1.5, zorder=8)
ax.add_patch(leg_bg)

ax.text(LGX+LGW/2, LGY-0.32,
    'CHU THICH', ha='center', va='center',
    fontsize=10, fontweight='bold', color='#1A202C', zorder=10)

legends = [
    (PK_BG,   '#975A16', '[PK]     Khoa chinh (Primary Key)'),
    (FK_BG,   '#2B6CB0', '[FK]     Khoa ngoai (Foreign Key)'),
    (PKFK_BG, '#276749', '[PK+FK]  Khoa chinh + Khoa ngoai'),
    (WHITE,   '#718096', '         Thuoc tinh thuong'),
]

for i, (bg, tc, text) in enumerate(legends):
    ly = LGY - 0.85 - i * 0.85
    rect = plt.Rectangle((LGX+0.18, ly-0.22), 0.60, 0.44,
        facecolor=bg, edgecolor='#CBD5E0', linewidth=1.0, zorder=9)
    ax.add_patch(rect)
    ax.text(LGX+0.95, ly, text,
        ha='left', va='center', fontsize=8.0, color='#1A202C', zorder=10)

# Moi quan he
ax.text(LGX+0.18, LGY-4.10,
    'Moi quan he:', ha='left', va='center',
    fontsize=9, fontweight='bold', color='#1A202C', zorder=10)

ax.plot([LGX+0.18, LGX+1.40], [LGY-4.60, LGY-4.60],
    color=LINE_C, linewidth=2, zorder=10)
ax.text(LGX+1.55, LGY-4.60, '-- Duong lien ket',
    ha='left', va='center', fontsize=8.0, color='#4A5568', zorder=10)

ax.text(LGX+0.18, LGY-5.15,
    '  1, N  = Ban so (Cardinality)',
    ha='left', va='center', fontsize=8.0, color='#4A5568', zorder=10,
    bbox=dict(boxstyle='round,pad=0.2', facecolor='white',
              edgecolor=LINE_C, linewidth=0.8))

# ═══════════════════════════════════════════════════════════════
# LUU ANH PNG
# ═══════════════════════════════════════════════════════════════
OUT_IMG  = r'D:\Do An Co So\car-website-dotnet\ERD_AutoHT.png'
OUT_DOCX = r'D:\Do An Co So\car-website-dotnet\ERD_AutoHT.docx'

# Thu luu vao duong dan goc (co dau tieng Viet)
try:
    import os
    base = r'D:\Đồ Án Cơ Sở\car-website-dotnet'
    OUT_IMG  = os.path.join(base, 'ERD_AutoHT.png')
    OUT_DOCX = os.path.join(base, 'ERD_AutoHT.docx')
except:
    pass

plt.tight_layout(pad=0.3)
plt.savefig(OUT_IMG, dpi=180, bbox_inches='tight',
    facecolor=fig.get_facecolor())
plt.close()
print('[OK] Anh ERD: ' + OUT_IMG)

# ═══════════════════════════════════════════════════════════════
# TAO FILE WORD  (kho A3 ngang)
# ═══════════════════════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width   = Cm(42.0)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin   = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ── Tieu de ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SO DO THUC THE - QUAN HE (ERD)')
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Du an AutoHT  --  He thong Quan ly va Mua ban Xe O to Truc tuyen')
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

doc.add_paragraph()

# ── Anh ERD ──
doc.add_picture(OUT_IMG, width=Cm(38.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Chu thich hinh ──
pn = doc.add_paragraph()
pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = pn.add_run('Hinh 1. So do ERD he thong AutoHT  |  Cong nghe: ASP.NET Core MVC + SQL Server')
rn.font.size = Pt(10); rn.italic = True
rn.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()

# ── Bang mo ta entities ──
h1 = doc.add_paragraph()
h1r = h1.add_run('MO TA CAC THUC THE (ENTITIES)')
h1r.bold = True; h1r.font.size = Pt(13)
h1r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

entity_info = [
    ('MauXe (Mau Xe)',
     'Thong tin tong quat ve moi dong/mau xe: ten, thuong hieu, gia niem yet, dong co.',
     'ModelID (PK)'),
    ('XeDonVi (Xe Don Vi)',
     'Moi chiec xe vat ly thuc te trong kho, dinh danh bang so VIN duy nhat.',
     'VIN (PK)'),
    ('DonHang (Don Hang)',
     'Ghi nhan giao dich mua hoac thue xe giua khach hang va mot xe cu the (VIN).',
     'OrderID (PK)'),
    ('KhachHang (Khach Hang)',
     'Thong tin ca nhan khach hang: ten, email, so dien thoai, dia chi.',
     'CustomerID (PK)'),
    ('PhieuDichVu (Phieu Dich Vu)',
     'Ghi nhan yeu cau bao duong/sua chua hau mai cua khach hang cho mot chiec xe.',
     'TicketID (PK)'),
    ('LinhKienPhuTung (Linh Kien)',
     'Danh muc linh kien/phu tung trong kho dich vu: ten, gia, so luong ton.',
     'PartID (PK)'),
    ('ChiTietPhieuDichVu (Chi Tiet)',
     'Bang trung gian giai quyet quan he N-N giua PhieuDichVu va LinhKienPhuTung. '
     'Luu so luong phu tung su dung cho moi phieu.',
     'TicketID + PartID (PK tong hop)'),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hcells = tbl.rows[0].cells
for i, h in enumerate(['Thuc the (Entity)', 'Mo ta', 'Khoa chinh (PK)']):
    hcells[i].text = h
    hcells[i].paragraphs[0].runs[0].bold = True
    hcells[i].paragraphs[0].runs[0].font.size = Pt(10)
    hcells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for en, desc, pk in entity_info:
    rc = tbl.add_row().cells
    rc[0].text = en; rc[0].paragraphs[0].runs[0].bold = True
    rc[1].text = desc
    rc[2].text = pk
    for c in rc: c.paragraphs[0].runs[0].font.size = Pt(9)

for row in tbl.rows:
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(23.0)
    row.cells[2].width = Cm(7.0)

doc.add_paragraph()

# ── Bang moi quan he ──
h2 = doc.add_paragraph()
h2r = h2.add_run('MOI QUAN HE GIUA CAC THUC THE (RELATIONSHIPS)')
h2r.bold = True; h2r.font.size = Pt(13)
h2r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

rel_info = [
    ('MauXe → XeDonVi',              '1 -- N',
     'Mot mau xe co nhieu xe don vi thuc te trong kho. (ModelID la FK trong XeDonVi)'),
    ('XeDonVi → DonHang',            '1 -- 1',
     'Mot xe don vi (VIN) chi nam tren toi da 1 don hang dang hoat dong. (VIN la FK trong DonHang)'),
    ('KhachHang → DonHang',          '1 -- N',
     'Mot khach hang co the dat nhieu don hang mua/thue. (CustomerID la FK trong DonHang)'),
    ('XeDonVi → PhieuDichVu',        '1 -- N',
     'Mot xe co the vao xuong nhieu lan (bao duong dinh ky, sua chua). (VIN la FK trong PhieuDichVu)'),
    ('KhachHang → PhieuDichVu',      '1 -- N',
     'Mot khach hang co the co nhieu phieu dich vu hau mai. (CustomerID la FK trong PhieuDichVu)'),
    ('PhieuDichVu <-> LinhKienPhuTung', 'N -- N',
     'Mot phieu dung nhieu phu tung; mot phu tung co the xuat hien o nhieu phieu. '
     'Giai quyet qua bang trung gian ChiTietPhieuDichVu.'),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
hc2 = tbl2.rows[0].cells
for i, h in enumerate(['Quan he', 'Ban so', 'Dien giai']):
    hc2[i].text = h
    hc2[i].paragraphs[0].runs[0].bold = True
    hc2[i].paragraphs[0].runs[0].font.size = Pt(10)
    hc2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for rel, card_s, desc in rel_info:
    rc = tbl2.add_row().cells
    rc[0].text = rel; rc[0].paragraphs[0].runs[0].bold = True
    rc[1].text = card_s; rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc[2].text = desc
    for c in rc: c.paragraphs[0].runs[0].font.size = Pt(9)

for row in tbl2.rows:
    row.cells[0].width = Cm(8.0)
    row.cells[1].width = Cm(3.0)
    row.cells[2].width = Cm(24.5)

doc.save(OUT_DOCX)
print('[OK] File Word: ' + OUT_DOCX)
print('[DONE] Hoan thanh!')
